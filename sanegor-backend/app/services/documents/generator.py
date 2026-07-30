"""Export of generated content to DOCX and PDF, right-to-left.

RTL output is the hard part of both formats:

* **DOCX** — python-docx has no RTL API, so the ``w:bidi`` (paragraph) and
  ``w:rtl`` (run) properties are written into the underlying XML directly.
* **PDF** — ReportLab does not reorder bidirectional text, so Hebrew has to be
  reversed manually per line before drawing, and a Hebrew-capable TrueType font
  must be registered (the built-in Type 1 fonts have no Hebrew glyphs).
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from app.core.logging import get_logger
from app.services.ai.prompts import DISCLAIMER_HE

logger = get_logger(__name__)

# Searched in order; the Docker image installs Noto Sans Hebrew.
_HEBREW_FONT_CANDIDATES = (
    "/usr/share/fonts/truetype/noto/NotoSansHebrew-Regular.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    "/System/Library/Fonts/Supplemental/Arial.ttf",
)

_HEBREW_BLOCK = re.compile(r"[֐-׿]")
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_MD_BULLET = re.compile(r"^[-*+]\s+(.*)$")
_MD_ORDERED = re.compile(r"^(\d+)[.)]\s+(.*)$")
_MD_BOLD = re.compile(r"\*\*(.+?)\*\*")
_MD_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


@dataclass(slots=True)
class ExportMetadata:
    """Header/footer information stamped onto an export."""

    title: str
    subtitle: str | None = None
    author: str = "סנגור"
    created: date | None = None
    include_disclaimer: bool = True


def strip_markdown(text: str) -> str:
    """Remove inline markdown emphasis, keeping the words."""
    text = _MD_BOLD.sub(r"\1", text)
    text = _MD_ITALIC.sub(r"\1", text)
    return text.replace("`", "")


def contains_hebrew(text: str) -> bool:
    return bool(_HEBREW_BLOCK.search(text))


# ---------------------------------------------------------------------- DOCX
class DocxExporter:
    """Renders markdown-ish content into a right-to-left Word document."""

    def render(self, body_markdown: str, meta: ExportMetadata) -> bytes:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        document = docx.Document()
        self._configure_styles(document, Pt)
        self._set_document_rtl(document)

        document.core_properties.title = meta.title
        document.core_properties.author = meta.author

        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(meta.title)
        run.bold = True
        run.font.size = Pt(20)
        self._set_run_rtl(run)

        if meta.subtitle:
            subtitle = document.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = subtitle.add_run(meta.subtitle)
            sub_run.font.size = Pt(11)
            sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._set_run_rtl(sub_run)

        document.add_paragraph()

        for line in body_markdown.splitlines():
            self._add_line(document, line, Pt, WD_ALIGN_PARAGRAPH)

        if meta.include_disclaimer:
            document.add_paragraph()
            note = document.add_paragraph()
            self._set_paragraph_rtl(note)
            note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            note_run = note.add_run(DISCLAIMER_HE)
            note_run.italic = True
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            self._set_run_rtl(note_run)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    # ------------------------------------------------------------- rendering
    def _add_line(self, document: object, line: str, Pt: object, align: object) -> None:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph()  # type: ignore[attr-defined]
            return

        if stripped in {"---", "***", "___"}:
            paragraph = document.add_paragraph()  # type: ignore[attr-defined]
            self._add_bottom_border(paragraph)
            return

        if (match := _MD_HEADING.match(stripped)) is not None:
            level = min(len(match.group(1)), 4)
            paragraph = document.add_paragraph()  # type: ignore[attr-defined]
            self._set_paragraph_rtl(paragraph)
            paragraph.alignment = align.RIGHT  # type: ignore[attr-defined]
            run = paragraph.add_run(strip_markdown(match.group(2)))
            run.bold = True
            run.font.size = Pt({1: 17, 2: 15, 3: 13, 4: 12}[level])  # type: ignore[operator]
            self._set_run_rtl(run)
            return

        text = stripped
        style: str | None = None
        if (bullet := _MD_BULLET.match(stripped)) is not None:
            text, style = bullet.group(1), "List Bullet"
        elif (ordered := _MD_ORDERED.match(stripped)) is not None:
            text, style = ordered.group(2), "List Number"

        paragraph = (
            document.add_paragraph(style=style)  # type: ignore[attr-defined]
            if style
            else document.add_paragraph()  # type: ignore[attr-defined]
        )
        self._set_paragraph_rtl(paragraph)
        paragraph.alignment = align.RIGHT  # type: ignore[attr-defined]
        self._add_formatted_runs(paragraph, text)

    def _add_formatted_runs(self, paragraph: object, text: str) -> None:
        """Split on ``**bold**`` so emphasis survives the export."""
        position = 0
        for match in _MD_BOLD.finditer(text):
            if match.start() > position:
                self._add_run(paragraph, text[position : match.start()], bold=False)
            self._add_run(paragraph, match.group(1), bold=True)
            position = match.end()
        if position < len(text):
            self._add_run(paragraph, strip_markdown(text[position:]), bold=False)

    def _add_run(self, paragraph: object, text: str, *, bold: bool) -> None:
        if not text:
            return
        run = paragraph.add_run(strip_markdown(text))  # type: ignore[attr-defined]
        run.bold = bold
        self._set_run_rtl(run)

    # ------------------------------------------------------------- XML plumbing
    @staticmethod
    def _qn(tag: str) -> str:
        from docx.oxml.ns import qn

        return qn(tag)

    @staticmethod
    def _configure_styles(document: object, Pt: object) -> None:
        style = document.styles["Normal"]  # type: ignore[index]
        style.font.name = "Arial"
        style.font.size = Pt(11)  # type: ignore[operator]
        style.paragraph_format.space_after = Pt(6)  # type: ignore[operator]
        style.paragraph_format.line_spacing = 1.15

    def _set_document_rtl(self, document: object) -> None:
        """Mark every section as right-to-left."""
        from docx.oxml import OxmlElement

        for section in document.sections:  # type: ignore[attr-defined]
            properties = section._sectPr
            bidi = OxmlElement("w:bidi")
            properties.append(bidi)

    def _set_paragraph_rtl(self, paragraph: object) -> None:
        from docx.oxml import OxmlElement

        properties = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(self._qn("w:val"), "1")
        properties.append(bidi)

    def _set_run_rtl(self, run: object) -> None:
        from docx.oxml import OxmlElement

        properties = run._element.get_or_add_rPr()
        rtl = OxmlElement("w:rtl")
        rtl.set(self._qn("w:val"), "1")
        properties.append(rtl)

    def _add_bottom_border(self, paragraph: object) -> None:
        from docx.oxml import OxmlElement

        properties = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(self._qn("w:val"), "single")
        bottom.set(self._qn("w:sz"), "6")
        bottom.set(self._qn("w:color"), "CCCCCC")
        borders.append(bottom)
        properties.append(borders)


# ----------------------------------------------------------------------- PDF
class PdfExporter:
    """Renders markdown-ish content into an RTL PDF."""

    _FONT_NAME = "SanegorHebrew"
    _MARGIN = 56
    _LEADING = 17

    def __init__(self) -> None:
        self._font_registered = False

    def render(self, body_markdown: str, meta: ExportMetadata) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        font = self._ensure_font()
        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        pdf.setTitle(meta.title)
        pdf.setAuthor(meta.author)

        width, height = A4
        right = width - self._MARGIN
        usable = width - 2 * self._MARGIN
        y = height - self._MARGIN

        pdf.setFont(font, 18)
        pdf.drawRightString(right, y, self._shape(meta.title))
        y -= 26

        if meta.subtitle:
            pdf.setFont(font, 10)
            pdf.setFillGray(0.4)
            pdf.drawRightString(right, y, self._shape(meta.subtitle))
            pdf.setFillGray(0)
            y -= 20

        pdf.setLineWidth(0.5)
        pdf.setStrokeGray(0.8)
        pdf.line(self._MARGIN, y, right, y)
        y -= 24

        for line in body_markdown.splitlines():
            y = self._draw_line(pdf, line, y, right, usable, font, height)

        if meta.include_disclaimer:
            if y < self._MARGIN + 60:
                pdf.showPage()
                y = height - self._MARGIN
            y -= 12
            pdf.setStrokeGray(0.85)
            pdf.line(self._MARGIN, y, right, y)
            y -= 16
            pdf.setFont(font, 8)
            pdf.setFillGray(0.45)
            for wrapped in self._wrap(pdf, DISCLAIMER_HE, usable, font, 8):
                pdf.drawRightString(right, y, self._shape(wrapped))
                y -= 11

        pdf.save()
        return buffer.getvalue()

    # ------------------------------------------------------------- rendering
    def _draw_line(
        self,
        pdf: object,
        line: str,
        y: float,
        right: float,
        usable: float,
        font: str,
        page_height: float,
    ) -> float:
        stripped = line.strip()
        if not stripped:
            return y - 8

        size, bullet = 11.0, ""
        if (heading := _MD_HEADING.match(stripped)) is not None:
            level = min(len(heading.group(1)), 4)
            size = {1: 16.0, 2: 14.0, 3: 12.5, 4: 11.5}[level]
            stripped = heading.group(2)
            y -= 6
        elif (item := _MD_BULLET.match(stripped)) is not None:
            stripped, bullet = item.group(1), "• "
        elif (ordered := _MD_ORDERED.match(stripped)) is not None:
            stripped, bullet = ordered.group(2), f"{ordered.group(1)}. "

        text = bullet + strip_markdown(stripped)
        pdf.setFont(font, size)  # type: ignore[attr-defined]
        pdf.setFillGray(0)  # type: ignore[attr-defined]

        for wrapped in self._wrap(pdf, text, usable, font, size):
            if y < self._MARGIN + 40:
                pdf.showPage()  # type: ignore[attr-defined]
                pdf.setFont(font, size)  # type: ignore[attr-defined]
                y = page_height - self._MARGIN
            pdf.drawRightString(right, y, self._shape(wrapped))  # type: ignore[attr-defined]
            y -= self._LEADING if size <= 11 else size + 7
        return y

    @staticmethod
    def _wrap(pdf: object, text: str, max_width: float, font: str, size: float) -> list[str]:
        """Greedy word wrap measured with the real font metrics."""
        words = text.split()
        if not words:
            return [""]
        lines: list[str] = []
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if pdf.stringWidth(candidate, font, size) <= max_width:  # type: ignore[attr-defined]
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines

    @staticmethod
    def _shape(text: str) -> str:
        """Reverse Hebrew runs so ReportLab draws them in visual order.

        ReportLab has no bidi engine. Splitting on script boundaries and
        reversing only the Hebrew runs keeps embedded Latin words and numbers
        (case numbers, amounts, dates) readable.
        """
        if not contains_hebrew(text):
            return text
        tokens = re.findall(r"[֐-׿\"'׳״]+|[^֐-׿]+", text)
        out: list[str] = []
        for token in tokens:
            out.append(token[::-1] if _HEBREW_BLOCK.search(token) else token)
        return "".join(reversed(out))

    # ---------------------------------------------------------------- fonts
    def _ensure_font(self) -> str:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        if self._font_registered:
            return self._FONT_NAME

        for candidate in _HEBREW_FONT_CANDIDATES:
            path = Path(candidate)
            if not path.exists():
                continue
            try:
                pdfmetrics.registerFont(TTFont(self._FONT_NAME, str(path)))
            except Exception as exc:
                logger.warning("pdf_font_failed", path=str(path), error=str(exc))
                continue
            self._font_registered = True
            logger.info("pdf_font_registered", path=str(path))
            return self._FONT_NAME

        # Helvetica has no Hebrew glyphs, but a Latin-only PDF beats a crash.
        logger.error(
            "pdf_hebrew_font_missing",
            detail="install fonts-noto-core; Hebrew will not render correctly",
        )
        return "Helvetica"
