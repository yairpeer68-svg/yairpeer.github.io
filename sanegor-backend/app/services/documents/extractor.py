"""Text extraction from uploaded files.

PDF, DOCX and plain text are parsed directly; images always go through OCR.
A PDF that yields almost no text is treated as a scan and re-processed with
OCR — which is the common case for Israeli court and municipal documents.

All parsing runs in a worker thread: pypdf and python-docx are synchronous and
CPU-bound, and blocking the event loop would stall every other request.
"""

from __future__ import annotations

import asyncio
import io
import re
import unicodedata
from dataclasses import dataclass, field
from typing import TYPE_CHECKING

from app.core.errors import UnsupportedMediaTypeError, ValidationError
from app.core.logging import get_logger

if TYPE_CHECKING:
    from app.services.documents.ocr import OcrService

logger = get_logger(__name__)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TXT_MIME = "text/plain"
IMAGE_MIMES = frozenset({"image/jpeg", "image/png", "image/jpg"})

# Below this many characters per page a PDF is almost certainly a scan.
_SCANNED_PDF_CHARS_PER_PAGE = 40

_HEBREW_RE = re.compile(r"[֐-׿]")
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_RTL_MARKS_RE = re.compile(r"[‎‏‪-‮⁦-⁩]")


@dataclass(slots=True)
class ExtractionResult:
    """Outcome of parsing one uploaded file."""

    text: str
    page_count: int | None = None
    word_count: int = 0
    language: str = "he"
    used_ocr: bool = False
    warnings: list[str] = field(default_factory=list)

    @property
    def is_empty(self) -> bool:
        return not self.text.strip()


def normalise_text(raw: str) -> str:
    """Clean extracted text for storage, display and embedding.

    Strips control characters and bidirectional marks (PDF extractors sprinkle
    them liberally through Hebrew text, and they corrupt both search and token
    counting), then normalises to NFC and collapses runaway blank lines.
    """
    text = _CONTROL_RE.sub("", raw)
    text = _RTL_MARKS_RE.sub("", text)
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.strip() for line in text.split("\n")).strip()


def detect_language(text: str) -> str:
    """Return ``he`` or ``en`` based on which script dominates."""
    sample = text[:4000]
    hebrew = len(_HEBREW_RE.findall(sample))
    latin = sum(1 for ch in sample if "a" <= ch.lower() <= "z")
    if hebrew == 0 and latin == 0:
        return "he"
    return "he" if hebrew >= latin else "en"


class DocumentExtractor:
    """Dispatches an upload to the right parser."""

    def __init__(self, ocr_service: OcrService | None = None) -> None:
        self._ocr = ocr_service

    async def extract(self, data: bytes, content_type: str, filename: str) -> ExtractionResult:
        """Extract text from ``data``.

        Raises:
            UnsupportedMediaTypeError: for a type we cannot parse.
            ValidationError: when the file is corrupt or yields nothing.
        """
        content_type = (content_type or "").split(";")[0].strip().lower()

        if content_type == PDF_MIME:
            result = await asyncio.to_thread(self._extract_pdf, data)
            if self._looks_scanned(result) and self._ocr is not None:
                logger.info("pdf_looks_scanned", filename=filename)
                ocr_result = await self._ocr.extract_from_pdf(data)
                if len(ocr_result.text) > len(result.text):
                    ocr_result.warnings.append("הטקסט חולץ באמצעות OCR")
                    return self._finalise(ocr_result)
            return self._finalise(result)

        if content_type == DOCX_MIME:
            return self._finalise(await asyncio.to_thread(self._extract_docx, data))

        if content_type == TXT_MIME:
            return self._finalise(await asyncio.to_thread(self._extract_txt, data))

        if content_type in IMAGE_MIMES:
            if self._ocr is None:
                raise UnsupportedMediaTypeError("OCR אינו מופעל בשרת זה")
            return self._finalise(await self._ocr.extract_from_image(data))

        raise UnsupportedMediaTypeError(
            "סוג הקובץ אינו נתמך", details={"content_type": content_type, "filename": filename}
        )

    @staticmethod
    def _looks_scanned(result: ExtractionResult) -> bool:
        pages = result.page_count or 1
        return len(result.text.strip()) < _SCANNED_PDF_CHARS_PER_PAGE * pages

    @staticmethod
    def _finalise(result: ExtractionResult) -> ExtractionResult:
        result.text = normalise_text(result.text)
        if result.is_empty:
            raise ValidationError(
                "לא הצלחנו לחלץ טקסט מהקובץ. נסה קובץ ברור יותר או קובץ טקסט"
            )
        result.word_count = len(result.text.split())
        result.language = detect_language(result.text)
        return result

    # ----------------------------------------------------------------- parsers
    @staticmethod
    def _extract_pdf(data: bytes) -> ExtractionResult:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError

        try:
            reader = PdfReader(io.BytesIO(data))
        except PdfReadError as exc:
            raise ValidationError("קובץ ה-PDF פגום או מוגן בסיסמה") from exc

        if reader.is_encrypted:
            try:
                reader.decrypt("")  # empty-owner-password PDFs are common
            except Exception as exc:  # noqa: BLE001
                raise ValidationError("קובץ ה-PDF מוגן בסיסמה") from exc

        warnings: list[str] = []
        parts: list[str] = []
        for number, page in enumerate(reader.pages, start=1):
            try:
                parts.append(page.extract_text() or "")
            except Exception as exc:  # noqa: BLE001 - skip the bad page, keep the rest
                logger.warning("pdf_page_failed", page=number, error=str(exc))
                warnings.append(f"עמוד {number} לא נקרא במלואו")

        return ExtractionResult(
            text="\n\n".join(parts), page_count=len(reader.pages), warnings=warnings
        )

    @staticmethod
    def _extract_docx(data: bytes) -> ExtractionResult:
        import docx
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = docx.Document(io.BytesIO(data))
        except (PackageNotFoundError, KeyError, ValueError) as exc:
            raise ValidationError("קובץ ה-Word פגום או אינו בפורמט DOCX") from exc

        parts = [p.text for p in document.paragraphs if p.text.strip()]
        # Tables carry the substance in many contracts; render them as pipes.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return ExtractionResult(text="\n".join(parts))

    @staticmethod
    def _extract_txt(data: bytes) -> ExtractionResult:
        for encoding in ("utf-8", "utf-8-sig", "windows-1255", "iso-8859-8"):
            try:
                return ExtractionResult(text=data.decode(encoding))
            except UnicodeDecodeError:
                continue
        return ExtractionResult(
            text=data.decode("utf-8", errors="replace"),
            warnings=["חלק מהתווים לא פוענחו"],
        )
